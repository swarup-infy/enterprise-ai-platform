"""
Document Extractor Service.

Extracts text from supported document formats:

- PDF
- DOCX
- TXT
- Markdown
- CSV
- XLSX
- XLS
- JSON
"""

from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any

import pandas as pd
import pymupdf
from docx import Document as DocxDocument


class ExtractorService:
    """
    Extract text from supported document formats.

    The extracted output is normalized into plain text suitable
    for downstream chunking, embedding, and RAG processing.
    """

    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".docx",
        ".txt",
        ".md",
        ".csv",
        ".xlsx",
        ".xls",
        ".json",
    }

    TEXT_ENCODINGS = (
        "utf-8",
        "utf-8-sig",
        "cp1252",
        "latin-1",
    )

    CSV_SAMPLE_SIZE = 8192

    # ==========================================================================
    # Public API
    # ==========================================================================

    def extract(
        self,
        file_path: str | Path,
    ) -> str:
        """
        Extract text from a supported document.

        Args:
            file_path: Path to the document.

        Returns:
            Cleaned extracted text.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file type is unsupported or contains
                no extractable text.
            RuntimeError: If extraction fails unexpectedly.
        """

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(
                f"Document file not found: {path}"
            )

        if not path.is_file():
            raise ValueError(
                f"Document path is not a file: {path}"
            )

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                "Unsupported file type: "
                f"{extension or 'unknown'}"
            )

        try:
            if extension == ".pdf":
                text = self._extract_pdf(path)

            elif extension == ".docx":
                text = self._extract_docx(path)

            elif extension in {".txt", ".md"}:
                text = self._extract_text(path)

            elif extension == ".csv":
                text = self._extract_csv(path)

            elif extension in {".xlsx", ".xls"}:
                text = self._extract_excel(path)

            elif extension == ".json":
                text = self._extract_json(path)

            else:
                raise ValueError(
                    f"Unsupported file type: {extension}"
                )

        except (FileNotFoundError, ValueError):
            raise

        except Exception as exc:
            raise RuntimeError(
                f"Failed to extract text from "
                f"'{path.name}': {exc}"
            ) from exc

        cleaned_text = self._clean_text(text)

        if not cleaned_text:
            raise ValueError(
                f"No extractable text found in "
                f"'{path.name}'. "
                "The document may be empty or image-only."
            )

        return cleaned_text

    # ==========================================================================
    # PDF
    # ==========================================================================

    @staticmethod
    def _extract_pdf(
        path: Path,
    ) -> str:
        """
        Extract text from all PDF pages.

        Page boundaries are preserved so downstream processing
        retains useful document structure.
        """

        pages: list[str] = []

        with pymupdf.open(path) as pdf:
            for page_number, page in enumerate(
                pdf,
                start=1,
            ):
                page_text = page.get_text(
                    "text"
                ).strip()

                if page_text:
                    pages.append(
                        f"Page {page_number}\n"
                        f"{page_text}"
                    )

        return "\n\n".join(pages)

    # ==========================================================================
    # DOCX
    # ==========================================================================

    @staticmethod
    def _extract_docx(
        path: Path,
    ) -> str:
        """
        Extract text from DOCX paragraphs and tables.

        Table boundaries are explicitly labelled so their content
        remains understandable after chunking.
        """

        document = DocxDocument(path)

        sections: list[str] = []

        # ----------------------------------------------------------------------
        # Paragraphs
        # ----------------------------------------------------------------------

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                sections.append(text)

        # ----------------------------------------------------------------------
        # Tables
        # ----------------------------------------------------------------------

        for table_index, table in enumerate(
            document.tables,
            start=1,
        ):
            sections.append(
                f"Table {table_index}"
            )

            for row in table.rows:
                cells = [
                    cell.text.strip()
                    for cell in row.cells
                ]

                if any(cells):
                    sections.append(
                        " | ".join(cells)
                    )

        return "\n".join(sections)

    # ==========================================================================
    # TXT / Markdown
    # ==========================================================================

    @classmethod
    def _extract_text(
        cls,
        path: Path,
    ) -> str:
        """
        Extract UTF-based or legacy-encoded text files.

        Several common encodings are attempted to support
        real-world uploaded documents.
        """

        last_error: UnicodeDecodeError | None = None

        for encoding in cls.TEXT_ENCODINGS:
            try:
                return path.read_text(
                    encoding=encoding,
                )

            except UnicodeDecodeError as exc:
                last_error = exc

        raise ValueError(
            f"Unable to decode text file: {path.name}"
        ) from last_error

    # ==========================================================================
    # CSV
    # ==========================================================================

    @classmethod
    def _extract_csv(
        cls,
        path: Path,
    ) -> str:
        """
        Extract CSV rows into readable text.

        The implementation attempts to detect the delimiter
        automatically and falls back to standard CSV parsing.
        """

        last_error: UnicodeDecodeError | None = None

        for encoding in cls.TEXT_ENCODINGS:
            try:
                with path.open(
                    "r",
                    newline="",
                    encoding=encoding,
                ) as file:
                    sample = file.read(
                        cls.CSV_SAMPLE_SIZE
                    )

                    file.seek(0)

                    try:
                        dialect = csv.Sniffer().sniff(
                            sample,
                        )
                    except csv.Error:
                        dialect = csv.excel

                    reader = csv.reader(
                        file,
                        dialect,
                    )

                    rows: list[str] = []

                    for row in reader:
                        cleaned_row = [
                            str(cell).strip()
                            for cell in row
                        ]

                        if any(cleaned_row):
                            rows.append(
                                " | ".join(
                                    cleaned_row
                                )
                            )

                    return "\n".join(rows)

            except UnicodeDecodeError as exc:
                last_error = exc

        raise ValueError(
            f"Unable to decode CSV file: {path.name}"
        ) from last_error

    # ==========================================================================
    # Excel
    # ==========================================================================

    @staticmethod
    def _extract_excel(
        path: Path,
    ) -> str:
        """
        Extract text from all Excel worksheets.

        Each worksheet is labelled explicitly.
        """

        workbook = pd.read_excel(
            path,
            sheet_name=None,
        )

        sections: list[str] = []

        for sheet_name, dataframe in workbook.items():
            sections.append(
                f"Sheet: {sheet_name}"
            )

            if dataframe.empty:
                sections.append(
                    "[Empty sheet]"
                )
                continue

            dataframe = dataframe.fillna("")

            table_text = dataframe.to_string(
                index=False,
            ).strip()

            if table_text:
                sections.append(
                    table_text
                )
            else:
                sections.append(
                    "[Empty sheet]"
                )

        return "\n\n".join(sections)

    # ==========================================================================
    # JSON
    # ==========================================================================

    @staticmethod
    def _extract_json(
        path: Path,
    ) -> str:
        """
        Extract JSON content and convert it into readable text.

        Structured JSON is pretty-printed rather than flattened so
        nested objects and arrays remain understandable to the LLM.
        """

        raw_text = path.read_text(
            encoding="utf-8-sig",
        )

        try:
            data: Any = json.loads(
                raw_text
            )

        except json.JSONDecodeError as exc:
            raise ValueError(
                f"Invalid JSON file: {path.name}. "
                f"{exc.msg} at line {exc.lineno}, "
                f"column {exc.colno}."
            ) from exc

        return json.dumps(
            data,
            indent=2,
            ensure_ascii=False,
        )

    # ==========================================================================
    # Text Cleaning
    # ==========================================================================

    @staticmethod
    def _clean_text(
        text: str,
    ) -> str:
        """
        Normalize extracted text.

        Removes trailing whitespace, collapses repeated blank
        lines, and preserves meaningful line boundaries.
        """

        if not text:
            return ""

        cleaned_lines: list[str] = []

        previous_blank = False

        for raw_line in text.splitlines():
            line = raw_line.strip()

            if not line:
                if not previous_blank:
                    cleaned_lines.append("")

                previous_blank = True
                continue

            cleaned_lines.append(line)
            previous_blank = False

        return "\n".join(
            cleaned_lines
        ).strip()


__all__ = [
    "ExtractorService",
]